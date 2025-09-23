import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from docx import Document
import re
from typing import List, Tuple, Optional

class WordFormFiller:
    def __init__(self, root):
        self.root = root
        self.root.title("Word 表格填寫工具")
        self.root.geometry("600x400")
        
        # 檔案路徑變數
        self.source_file = tk.StringVar()
        self.target_file = tk.StringVar()
        
        self.setup_ui()
    
    def setup_ui(self):
        # 主框架
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 標題
        title_label = ttk.Label(main_frame, text="Word 表格填寫工具", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # 解析卷檔案選擇
        ttk.Label(main_frame, text="解析卷檔案:").grid(row=1, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.source_file, width=50).grid(row=1, column=1, padx=5, pady=5)
        ttk.Button(main_frame, text="選擇檔案", 
                  command=self.select_source_file).grid(row=1, column=2, pady=5)
        
        # 解答卷檔案選擇
        ttk.Label(main_frame, text="解答卷檔案:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ttk.Entry(main_frame, textvariable=self.target_file, width=50).grid(row=2, column=1, padx=5, pady=5)
        ttk.Button(main_frame, text="選擇檔案", 
                  command=self.select_target_file).grid(row=2, column=2, pady=5)
        
        # 處理按鈕
        process_btn = ttk.Button(main_frame, text="開始處理", 
                               command=self.process_files, style="Accent.TButton")
        process_btn.grid(row=3, column=0, columnspan=3, pady=20)
        
        # 進度條
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        
        # 狀態標籤
        self.status_label = ttk.Label(main_frame, text="請選擇檔案後點擊開始處理")
        self.status_label.grid(row=5, column=0, columnspan=3, pady=10)
        
        # 日誌區域
        log_frame = ttk.LabelFrame(main_frame, text="處理日誌", padding="10")
        log_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=10)
        
        self.log_text = tk.Text(log_frame, height=10, width=70)
        scrollbar = ttk.Scrollbar(log_frame, orient="vertical", command=self.log_text.yview)
        self.log_text.configure(yscrollcommand=scrollbar.set)
        
        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # 配置網格權重
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(6, weight=1)
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
    
    def log_message(self, message):
        self.log_text.insert(tk.END, f"{message}\n")
        self.log_text.see(tk.END)
        self.root.update()
    
    def select_source_file(self):
        filename = filedialog.askopenfilename(
            title="選擇解析卷檔案",
            filetypes=[("Word 文檔", "*.docx"), ("所有檔案", "*.*")]
        )
        if filename:
            self.source_file.set(filename)
            self.log_message(f"已選擇解析卷檔案: {os.path.basename(filename)}")
    
    def select_target_file(self):
        filename = filedialog.askopenfilename(
            title="選擇解答卷檔案",
            filetypes=[("Word 文檔", "*.docx"), ("所有檔案", "*.*")]
        )
        if filename:
            self.target_file.set(filename)
            self.log_message(f"已選擇解答卷檔案: {os.path.basename(filename)}")
    
    def parse_source_document(self, doc_path: str) -> List[Tuple[str, str, str]]:
        try:
            doc = Document(doc_path)
            questions = []
            
            self.log_message("正在解析源文檔...")
            
            full_text = ""
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text += text + "\n"
            
            questions = self._parse_questions_from_text(full_text)
            
            self.log_message(f"成功解析 {len(questions)} 個題目")
            return questions
            
        except Exception as e:
            self.log_message(f"解析源文檔時發生錯誤: {str(e)}")
            return []
    
    def _parse_questions_from_text(self, text: str) -> List[Tuple[str, str, str]]:
        questions = []
        lines = text.split('\n')
        
        current_question_num = 1
        current_answer = None
        current_explanation = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 識別答案行：支援單選和多選格式
            # 單選：答案 : (A) 或 答案 : (Ｃ)
            # 多選：答案 : (1)(A); (2)(D) 或 答案 : (1)(Ｂ); (2)(Ａ)
            if re.match(r'^答案\s*[：:]\s*[（(]', line):
                # 如果有前一個題目，先保存
                if current_answer:
                    questions.append((
                        f"{current_question_num}.",
                        current_answer,
                        current_explanation or ""
                    ))
                    self.log_message(f"解析題目 {current_question_num}.: 答案={current_answer}, 解析={'有' if current_explanation else '無'}")
                    current_question_num += 1
                current_answer = line
                current_explanation = None
            elif re.match(r'^解析：', line) and current_answer:
                current_explanation = line
            elif current_explanation and current_answer and line:
                current_explanation += " " + line
        
        if current_answer:
            questions.append((
                f"{current_question_num}.",
                current_answer,
                current_explanation or ""
            ))
            self.log_message(f"解析題目 {current_question_num}.: 答案={current_answer}, 解析={'有' if current_explanation else '無'}")
        
        return questions

    def _parse_with_fallback(self, text: str) -> List[Tuple[str, str, str]]:
        questions = []
        lines = text.split('\n')
        
        current_question = None
        current_answer = None
        current_explanation = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if re.match(r'^\d+\.', line):
                if current_question:
                    questions.append((
                        current_question, 
                        current_answer or "", 
                        current_explanation or ""
                    ))
                current_question = line
                current_answer = None
                current_explanation = None
            
            elif re.search(r'答案\s*:\s*\([A-D]\)', line) and current_question:
                current_answer = line
            elif re.search(r'解析\s*:', line) and current_question:
                current_explanation = line
            elif current_explanation and current_question:
                current_explanation += " " + line
        if current_question:
            questions.append((
                current_question, 
                current_answer or "", 
                current_explanation or ""
            ))
        
        return questions
    
    def fill_target_document(self, target_path: str, questions: List[Tuple[str, str, str]]):
        try:
            doc = Document(target_path)
            
            self.log_message("正在填寫目標文檔...")
            
            # 尋找表格
            tables = doc.tables
            if not tables:
                self.log_message("警告: 目標文檔中沒有找到表格")
                return
            table = tables[0]
            self.log_message(f"找到表格，共 {len(table.rows)} 行，{len(table.columns)} 列")
            required_rows = len(questions) + 1  # +1 為標題行
            if required_rows > len(table.rows):
                self.log_message(f"需要 {required_rows} 行，但表格只有 {len(table.rows)} 行，正在擴展表格...")
                for _ in range(required_rows - len(table.rows)):
                    new_row = table.add_row()
                    for cell in new_row.cells:
                        cell.text = ""
                self.log_message(f"表格已擴展到 {len(table.rows)} 行")
            filled_count = 0
            for i, (question, answer, explanation) in enumerate(questions):
                row_idx = i + 1  # 跳過標題行
                row = table.rows[row_idx]
                if len(row.cells) > 0:
                    row.cells[0].text = question
                    self.log_message(f"填寫題序: {question}")
                if len(row.cells) > 1:
                    clean_answer = answer.replace('答案：', '').replace('答案:', '').strip()
                    row.cells[1].text = clean_answer
                    self.log_message(f"填寫答案: {clean_answer}")
                if len(row.cells) > 2:
                    if explanation and explanation.strip():
                        clean_explanation = explanation.replace('解析：', '').replace('解析:', '').strip()
                        row.cells[2].text = clean_explanation
                        self.log_message(f"填寫解析: {clean_explanation[:50]}{'...' if len(clean_explanation) > 50 else ''}")
                    else:
                        row.cells[2].text = ""
                        self.log_message(f"題目 {question} 無解析，跳過解析欄位")
                
                filled_count += 1
            
            output_path = target_path.replace('.docx', '_已填寫.docx')
            doc.save(output_path)
            
            self.log_message(f"成功填寫 {filled_count} 個題目")
            self.log_message(f"已保存到: {os.path.basename(output_path)}")
            
        except Exception as e:
            self.log_message(f"填寫目標文檔時發生錯誤: {str(e)}")
    
    def process_files(self):
        if not self.source_file.get() or not self.target_file.get():
            messagebox.showerror("錯誤", "請先選擇解析卷和解答卷檔案")
            return
        
        if not os.path.exists(self.source_file.get()):
            messagebox.showerror("錯誤", "解析卷檔案不存在")
            return
        
        if not os.path.exists(self.target_file.get()):
            messagebox.showerror("錯誤", "解答卷檔案不存在")
            return
        
        self.progress.start()
        self.status_label.config(text="正在處理...")
        self.log_text.delete(1.0, tk.END)
        
        try:
            questions = self.parse_source_document(self.source_file.get())
            
            if not questions:
                messagebox.showwarning("警告", "未能從解析卷中提取到任何題目")
                return
            
            self.fill_target_document(self.target_file.get(), questions)
            
            self.status_label.config(text="處理完成！")
            messagebox.showinfo("完成", "檔案處理完成！請檢查輸出的檔案。")
            
        except Exception as e:
            self.log_message(f"處理過程中發生錯誤: {str(e)}")
            messagebox.showerror("錯誤", f"處理失敗: {str(e)}")
        finally:
            self.progress.stop()
            self.status_label.config(text="處理完成")

def main():
    root = tk.Tk()
    app = WordFormFiller(root)
    root.mainloop()

if __name__ == "__main__":
    main()
